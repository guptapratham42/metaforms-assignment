import io
import json
import logging
from typing import Any, cast

import jsonschema
from fastapi import HTTPException, UploadFile
from openai import OpenAI, OpenAIError
from openai.types.chat import ChatCompletionMessageParam

from models import ExtractResponse
from utils import (
    build_adaptive_prompt,
    estimate_complexity,
    extract_json_from_response,
    get_model_for_complexity,
)

import pandas as pd
import PyPDF2
from docx import Document
from docx.document import Document as DocumentType


logger = logging.getLogger(__name__)


async def read_and_validate_schema(schema_file: UploadFile) -> dict[str, Any]:
    """Read and validate JSON schema from uploaded file."""
    try:
        schema_bytes = await schema_file.read()
        schema = json.loads(schema_bytes.decode("utf-8"))
        jsonschema.Draft7Validator.check_schema(schema)
        logger.info("Schema loaded and validated successfully")
        return schema
    except jsonschema.exceptions.SchemaError as e:
        logger.error("Schema validation error: %s", e)
        raise HTTPException(status_code=400, detail=f"Invalid schema: {e}") from e
    except Exception as e:
        logger.error("Error reading schema file: %s", e)
        raise HTTPException(
            status_code=400, detail=f"Failed to read/validate schema: {e}"
        ) from e


async def read_document_file(document_file: UploadFile) -> str:
    """Read document content from uploaded file with multi-format support."""
    try:
        filename = document_file.filename.lower() if document_file.filename else ""
        content_type = document_file.content_type or ""

        logger.info("Processing file: %s (type: %s)", filename, content_type)

        if filename.endswith(".pdf") or "pdf" in content_type:
            return await _extract_pdf_text(document_file)
        elif filename.endswith(".csv") or "csv" in content_type:
            return await _process_csv_file(document_file)
        elif filename.endswith(".docx") or "wordprocessingml" in content_type:
            return await _extract_docx_text(document_file)
        elif filename.endswith((".doc", ".rtf")):
            logger.warning(
                "Legacy document format detected: %s. Please convert to .docx or .txt",
                filename,
            )
            raise HTTPException(
                status_code=400,
                detail="Legacy document formats (.doc, .rtf) are not supported. Please convert to .docx or .txt",
            )
        else:
            # Fallback to text processing for .txt, .md, etc.
            return await _read_plain_text_file(document_file)

    except HTTPException:
        raise
    except Exception as e:
        logger.error("Error reading document file %s: %s", filename, e)
        raise HTTPException(
            status_code=400, detail=f"Failed to read document file: {e}"
        ) from e


async def _read_plain_text_file(text_file: UploadFile) -> str:
    """Read plain text content from uploaded file."""
    text_bytes = await text_file.read()

    # Try different encodings if UTF-8 fails
    encodings = ["utf-8", "utf-8-sig", "latin1", "cp1252"]

    for encoding in encodings:
        try:
            text = text_bytes.decode(encoding)
            logger.info(
                "Text file loaded successfully with %s encoding, length=%d chars",
                encoding,
                len(text),
            )
            return text
        except UnicodeDecodeError:
            continue

    raise ValueError("Unable to decode text file with any supported encoding")


async def _extract_pdf_text(pdf_file: UploadFile) -> str:
    """Extract text content from PDF file."""
    pdf_bytes = await pdf_file.read()
    pdf_stream = io.BytesIO(pdf_bytes)

    try:
        pdf_reader = PyPDF2.PdfReader(pdf_stream)

        if len(pdf_reader.pages) == 0:
            raise ValueError("PDF file contains no pages")

        text_content = []
        for page_num, page in enumerate(pdf_reader.pages):
            try:
                page_text = page.extract_text()
                if page_text.strip():  # Only add non-empty pages
                    text_content.append(f"--- Page {page_num + 1} ---\n{page_text}")
            except Exception as e:
                logger.warning(
                    "Failed to extract text from page %d: %s", page_num + 1, e
                )
                continue

        if not text_content:
            raise ValueError("No readable text found in PDF")

        full_text = "\n\n".join(text_content)
        logger.info(
            "PDF extracted successfully, %d pages, %d chars",
            len(pdf_reader.pages),
            len(full_text),
        )
        return full_text

    except Exception as e:
        logger.error("PDF extraction failed: %s", e)
        raise ValueError(f"Failed to extract text from PDF: {e}")


def _extract_paragraph_text(doc: DocumentType) -> list[str]:
    """Extract text from all paragraphs in the document."""
    paragraph_texts = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            paragraph_texts.append(paragraph.text.strip())
    return paragraph_texts


def _extract_table_row_text(row) -> str:
    """Extract text from a single table row."""
    row_text = []
    for cell in row.cells:
        cell_text = cell.text.strip()
        if cell_text:
            row_text.append(cell_text)
    return " | ".join(row_text) if row_text else ""


def _extract_single_table_text(table) -> list[str]:
    """Extract text from a single table."""
    table_text = []
    for row in table.rows:
        row_text = _extract_table_row_text(row)
        if row_text:
            table_text.append(row_text)
    return table_text


def _extract_all_tables_text(doc: DocumentType) -> list[str]:
    """Extract text from all tables in the document."""
    all_table_content = []
    for table in doc.tables:
        table_text = _extract_single_table_text(table)
        if table_text:
            formatted_table = "\n--- Table ---\n" + "\n".join(table_text)
            all_table_content.append(formatted_table)
    return all_table_content


def _combine_extracted_content(
    paragraph_texts: list[str], table_texts: list[str]
) -> str:
    """Combine paragraph and table content into final text."""
    text_content = paragraph_texts + table_texts
    if not text_content:
        raise ValueError("No readable text found in DOCX")
    return "\n\n".join(text_content)


async def _extract_docx_text(docx_file: UploadFile) -> str:
    """Extract text content from DOCX file."""
    try:
        # Read and prepare document
        docx_bytes = await docx_file.read()
        docx_stream = io.BytesIO(docx_bytes)
        doc = Document(docx_stream)

        # Extract content using helper functions
        paragraph_texts = _extract_paragraph_text(doc)
        table_texts = _extract_all_tables_text(doc)

        # Combine and return result
        full_text = _combine_extracted_content(paragraph_texts, table_texts)
        logger.info("DOCX extracted successfully, %d chars", len(full_text))
        return full_text

    except Exception as e:
        logger.error("DOCX extraction failed: %s", e)
        raise ValueError(f"Failed to extract text from DOCX: {e}")


async def _process_csv_file(csv_file: UploadFile) -> str:
    """Process CSV file and convert to structured text format."""
    csv_bytes = await csv_file.read()
    csv_text = csv_bytes.decode("utf-8-sig")  # Handle BOM if present
    csv_stream = io.StringIO(csv_text)

    try:
        # Try to detect delimiter
        sample = csv_text[:1024]
        delimiter = _detect_csv_delimiter(sample)

        # Read CSV with pandas for better handling of complex cases
        df = pd.read_csv(csv_stream, delimiter=delimiter)

        if df.empty:
            raise ValueError("CSV file is empty")

        # Convert to structured text format
        text_parts = [
            "CSV Data Summary:",
            f"Columns ({len(df.columns)}): {', '.join(df.columns)}",
            f"Rows: {len(df)}",
            "",
            "Data Preview:",
        ]

        # Add column information
        for col in df.columns:
            col_info = f"Column '{col}': {df[col].dtype}"
            if df[col].dtype == "object":
                unique_count = df[col].nunique()
                col_info += f", {unique_count} unique values"
                if unique_count <= 10:
                    sample_values = df[col].dropna().unique()[:5]
                    col_info += f", samples: {list(sample_values)}"
            text_parts.append(col_info)

        text_parts.append("\nFirst 10 rows:")
        text_parts.append(df.head(10).to_string(index=False))

        if len(df) > 10:
            text_parts.append(f"\n... and {len(df) - 10} more rows")

        full_text = "\n".join(text_parts)
        logger.info(
            "CSV processed successfully, %d rows x %d cols, %d chars",
            len(df),
            len(df.columns),
            len(full_text),
        )
        return full_text

    except Exception as e:
        logger.error("CSV processing failed: %s", e)
        raise ValueError(f"Failed to process CSV file: {e}")


def _detect_csv_delimiter(sample: str) -> str:
    """Detect CSV delimiter from sample text."""
    common_delimiters = [",", ";", "\t", "|"]
    delimiter_counts = {}

    for delimiter in common_delimiters:
        count = sample.count(delimiter)
        if count > 0:
            delimiter_counts[delimiter] = count

    if delimiter_counts:
        return max(delimiter_counts, key=lambda x: delimiter_counts[x])

    return ","


def analyze_extraction_context(
    schema: dict[str, Any], text: str
) -> tuple[dict[str, int], str, list[dict[str, str]]]:
    """Analyze complexity and prepare extraction context."""
    complexity = estimate_complexity(schema)
    model = get_model_for_complexity(complexity, len(text))
    messages = build_adaptive_prompt(schema, text, complexity)

    logger.info(
        "Analysis: nested_objects=%d, max_depth=%d, literals=%d, using model=%s",
        complexity["nested_objects"],
        complexity["max_depth"],
        complexity["literals"],
        model,
    )

    return complexity, model, messages


def call_openai_with_retry(
    client: OpenAI,
    model: str,
    messages: list[dict[str, str]],
    complexity: dict[str, int],
    max_retries: int = 2,
) -> str:
    """Call OpenAI API with error handling and retries."""
    for attempt in range(max_retries + 1):
        try:
            completion = client.chat.completions.create(
                model=model,
                messages=cast(list[ChatCompletionMessageParam], messages),
                temperature=0,
                max_tokens=(16384 if complexity["nested_objects"] > 50 else 8192),
            )
            break
        except OpenAIError as e:
            if attempt == max_retries:
                logger.error("OpenAI API error after %d retries: %s", max_retries, e)
                raise HTTPException(status_code=502, detail="OpenAI API error") from e
            logger.warning(
                "OpenAI API error (attempt %d): %s, retrying...",
                attempt + 1,
                e,
            )
        except Exception as e:
            logger.error("Unexpected error calling OpenAI: %s", e)
            raise HTTPException(
                status_code=502, detail="Error calling OpenAI API"
            ) from e

    content = getattr(completion.choices[0].message, "content", None)
    if not content or not isinstance(content, str):
        logger.error("Invalid response content type: %s", type(content))
        raise HTTPException(status_code=502, detail="Invalid response from model")

    logger.info("Received response from OpenAI, length=%d chars", len(content))
    return content


def parse_model_response(content: str) -> dict[str, Any]:
    """Parse and extract JSON from model response with fallback handling."""
    json_content = extract_json_from_response(content)
    logger.info("Extracted JSON content, length=%d chars", len(json_content))

    try:
        result = json.loads(json_content)
        logger.info("Model output parsed as JSON successfully")
        return result
    except json.JSONDecodeError as e:
        logger.error(
            "JSON parsing error: %s; extracted content: %s",
            e,
            json_content[:500],
        )
        # Try one more time with just the original content
        try:
            result = json.loads(content)
            logger.info("Fallback JSON parsing succeeded")
            return result
        except json.JSONDecodeError as e2:
            logger.error("Fallback JSON parsing also failed: %s", e2)
            raise HTTPException(
                status_code=500,
                detail=f"Failed to parse model output as JSON. Content preview: {json_content[:200]}...",
            ) from e2


def validate_against_schema(
    result: dict[str, Any], schema: dict[str, Any]
) -> tuple[dict[str, str], dict[str, float]]:
    """Validate extraction result against schema and calculate confidence scores."""
    validator = jsonschema.Draft7Validator(schema)
    errors: dict[str, str] = {}
    confidence_scores: dict[str, float] = {}

    # Collect validation errors
    for error in validator.iter_errors(result):
        path = ".".join(str(p) for p in error.path)
        field_path = path or "root"
        errors[field_path] = error.message
        # Low confidence for validation errors
        confidence_scores[field_path] = 0.0

    # Calculate confidence scores for valid fields
    _calculate_confidence_scores(result, confidence_scores)

    logger.info("Schema validation completed, errors found: %d", len(errors))
    return errors, confidence_scores


def _calculate_confidence_scores(
    obj: Any, confidence_scores: dict[str, float], path: str = ""
) -> None:
    """Calculate confidence scores for extracted fields (simplified heuristic)."""
    if isinstance(obj, dict):
        _process_dict_object(obj, confidence_scores, path)
    elif isinstance(obj, list):
        _process_list_object(obj, confidence_scores, path)


def _process_dict_object(
    obj: dict[str, Any], confidence_scores: dict[str, float], path: str
) -> None:
    """Process dictionary objects for confidence score calculation."""
    for key, value in obj.items():
        current_path = f"{path}.{key}" if path else key

        if current_path not in confidence_scores:
            confidence_scores[current_path] = _get_value_confidence(value)

        if isinstance(value, (dict, list)):
            _calculate_confidence_scores(value, confidence_scores, current_path)


def _process_list_object(
    obj: list[Any], confidence_scores: dict[str, float], path: str
) -> None:
    """Process list objects for confidence score calculation."""
    for i, item in enumerate(obj):
        _calculate_confidence_scores(item, confidence_scores, f"{path}[{i}]")


def _get_value_confidence(value: Any) -> float:
    """Determine confidence score based on value type and content."""
    if value is None:
        return 0.3  # Low confidence for null values
    elif isinstance(value, str) and len(value.strip()) == 0:
        return 0.4  # Low confidence for empty strings
    else:
        return 0.8  # Default confidence


def compile_extraction_stats(
    complexity: dict[str, int],
    model: str,
    text_length: int,
    result: dict[str, Any],
    errors: dict[str, str],
    confidence_scores: dict[str, float],
) -> dict[str, Any]:
    """Compile comprehensive extraction statistics."""
    return {
        "complexity": complexity,
        "model_used": model,
        "input_length": text_length,
        "output_length": len(json.dumps(result)),
        "validation_errors": len(errors),
        "total_fields_processed": len(confidence_scores),
    }


async def perform_extraction(
    client: OpenAI,
    schema_file: UploadFile,
    text_file: UploadFile,
) -> ExtractResponse:
    """
    Main extraction orchestrator function with reduced cognitive complexity.
    Coordinates all extraction steps using smaller, focused functions.
    """
    logger.info(
        "Received extraction request with files: %s, %s",
        schema_file.filename,
        text_file.filename,
    )

    # Step 1: Read and validate inputs

    schema = await read_and_validate_schema(schema_file)
    text = await read_document_file(text_file)

    # Step 2: Analyze and prepare extraction context
    complexity, model, messages = analyze_extraction_context(schema, text)

    # Step 3: Call OpenAI API with retry logic
    content = call_openai_with_retry(client, model, messages, complexity)

    # Step 4: Parse the model response
    result = parse_model_response(content)

    # Step 5: Validate against schema and calculate confidence
    errors, confidence_scores = validate_against_schema(result, schema)

    # Step 6: Compile final statistics
    extraction_stats = compile_extraction_stats(
        complexity, model, len(text), result, errors, confidence_scores
    )

    return ExtractResponse(
        data=result,
        flagged_fields=confidence_scores,
        extraction_stats=extraction_stats,
    )
